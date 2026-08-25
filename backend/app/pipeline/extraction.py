"""Text/visual extraction: PDF text layer, OCR fallback, DOCX and image OCR."""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_SUFFIXES = {".pdf", ".png", ".jpg", ".jpeg", ".docx", ".txt"}


class ExtractionError(RuntimeError):
    pass


def extract_text(path: Path) -> tuple[str, int, float]:
    """Return (text, page_count, ocr_quality) for a stored document."""
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ExtractionError(f"Unsupported file type: {suffix}")
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore"), 1, 1.0
    return _ocr_image(path)


def _extract_pdf(path: Path) -> tuple[str, int, float]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages).strip()
    if len(text) >= 40:
        return text, len(reader.pages), 0.98
    ocr_text = _ocr_pdf(path)
    if ocr_text.strip():
        return ocr_text, len(reader.pages), 0.82
    return text, len(reader.pages), 0.4


def _ocr_pdf(path: Path) -> str:
    try:
        import pytesseract
        from pdf2image import convert_from_path

        images = convert_from_path(str(path), dpi=200)
        return "\n".join(pytesseract.image_to_string(image) for image in images)
    except Exception as exc:  # noqa: BLE001 - OCR is best-effort
        logger.warning("PDF OCR failed for %s: %s", path.name, exc)
        return ""


def _ocr_image(path: Path) -> tuple[str, int, float]:
    try:
        import pytesseract
        from PIL import Image

        with Image.open(path) as image:
            data = pytesseract.image_to_data(
                image, output_type=pytesseract.Output.DICT
            )
        lines: dict[tuple[int, int, int], list[str]] = {}
        confs: list[int] = []
        for index, word in enumerate(data["text"]):
            if not word.strip():
                continue
            key = (data["block_num"][index], data["par_num"][index], data["line_num"][index])
            lines.setdefault(key, []).append(word)
            confidence = int(data["conf"][index])
            if confidence >= 0:
                confs.append(confidence)
        text = "\n".join(" ".join(words) for _, words in sorted(lines.items()))
        quality = (sum(confs) / len(confs) / 100) if confs else 0.5
        return text, 1, round(quality, 4)
    except Exception as exc:  # noqa: BLE001 - OCR is best-effort
        logger.warning("Image OCR failed for %s: %s", path.name, exc)
        return "", 1, 0.35


def _extract_docx(path: Path) -> tuple[str, int, float]:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip(), 1, 0.97
