"""IBM Docling document parser for structured document extraction."""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_SUFFIXES = {".pdf", ".png", ".jpg", ".jpeg", ".docx", ".txt"}


class DoclingParserError(RuntimeError):
    pass


def parse_document(path: Path) -> tuple[str, int, float]:
    """
    Parse document using IBM Docling DocumentConverter.
    
    Returns:
        tuple: (markdown_text, page_count, quality_score)
    """
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise DoclingParserError(f"Unsupported file type: {suffix}")
    
    try:
        from docling.document_converter import DocumentConverter
        from docling.datamodel.base_models import InputFormat
        
        # Try to configure Docling options for optimized performance
        # Disable OCR for native PDFs to speed up processing
        try:
            from docling.document_converter import DocumentConverterOptions
            options = DocumentConverterOptions(
                do_ocr=(suffix != ".pdf"),  # Disable OCR for PDFs, enable for images
                do_table_structure=True,    # Retain table parsing
                do_document_structure=True  # Retain layout parsing
            )
            converter = DocumentConverter(options=options)
            logger.info(f"Docling configured with optimized options for {suffix}")
        except ImportError:
            # Fallback to default converter if options not available
            converter = DocumentConverter()
            logger.info(f"Docling using default configuration for {suffix}")
        
        # Convert document to Docling document
        result = converter.convert(path)
        
        # Export to markdown for clean text extraction
        markdown_text = result.document.export_to_markdown()
        
        # Get page count
        page_count = len(result.document.pages) if hasattr(result.document, 'pages') else 1
        
        # Calculate quality score based on content richness
        quality_score = _calculate_quality_score(markdown_text, page_count)
        
        logger.info(f"Docling parsed {path.name}: {len(markdown_text)} chars, {page_count} pages, quality={quality_score}")
        
        return markdown_text, page_count, quality_score
        
    except ImportError:
        logger.warning("Docling not installed, falling back to basic extraction")
        return _fallback_extraction(path)
    except Exception as exc:
        logger.error(f"Docling parsing failed for {path.name}: {exc}")
        raise DoclingParserError(f"Docling parsing failed: {exc}") from exc


def _calculate_quality_score(markdown_text: str, page_count: int) -> float:
    """Calculate quality score based on extracted content."""
    if not markdown_text or len(markdown_text.strip()) < 10:
        return 0.3
    
    # Base score from content length
    length_score = min(1.0, len(markdown_text) / 5000)
    
    # Bonus for structured elements (headings, tables, lists)
    structure_bonus = 0.0
    if "##" in markdown_text:  # Has headings
        structure_bonus += 0.1
    if "|" in markdown_text and "---" in markdown_text:  # Has tables
        structure_bonus += 0.15
    if "-" in markdown_text and markdown_text.count("-") > 5:  # Has lists
        structure_bonus += 0.1
    
    # Normalize page count influence
    page_factor = min(1.0, page_count / 10)
    
    quality = 0.5 + (0.3 * length_score) + structure_bonus + (0.1 * page_factor)
    return round(min(0.99, quality), 4)


def _fallback_extraction(path: Path) -> tuple[str, int, float]:
    """Fallback extraction when Docling is not available."""
    suffix = path.suffix.lower()
    
    if suffix == ".txt":
        text = path.read_text(encoding="utf-8", errors="ignore")
        return text, 1, 0.8
    
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages).strip()
            return text, len(reader.pages), 0.7
        except Exception as exc:
            logger.warning(f"Fallback PDF extraction failed: {exc}")
            return "", 1, 0.3
    
    if suffix == ".docx":
        try:
            import docx
            document = docx.Document(str(path))
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(parts).strip()
            return text, 1, 0.75
        except Exception as exc:
            logger.warning(f"Fallback DOCX extraction failed: {exc}")
            return "", 1, 0.3
    
    # For images, return empty with low quality
    return "", 1, 0.2


def extract_lineage_context(markdown_text: str, field_key: str) -> str:
    """
    Extract relevant context snippet for a field from the markdown text.
    This provides lineage tracking for governance purposes.
    """
    if not markdown_text:
        return ""
    
    lines = markdown_text.split('\n')
    
    # Simple heuristic: find lines that might contain the field
    key_lower = field_key.lower().replace('_', ' ')
    
    for i, line in enumerate(lines):
        if key_lower in line.lower():
            # Return context around the match (3 lines before and after)
            start = max(0, i - 3)
            end = min(len(lines), i + 4)
            context = '\n'.join(lines[start:end])
            return context[:500]  # Limit context length
    
    # If no direct match, return first few lines as general context
    return '\n'.join(lines[:5])[:500]